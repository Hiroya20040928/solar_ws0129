from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "outputs" / "related_paper_review_202607"
SOURCE_DOCX = REPORT_DIR / "関連論文調査報告書_B4小林_2026_0730_原著再調査版.docx"
OUTPUT_DOCX = REPORT_DIR / "関連論文調査報告書_B4小林_2026_0731_最終選定版.docx"
FIGURE_DIR = REPORT_DIR / "source_figures_used_in_report"


def copy_paragraph_format(source, target):
    if target._p.pPr is not None:
        target._p.remove(target._p.pPr)
    if source._p.pPr is not None:
        target._p.insert(0, deepcopy(source._p.pPr))


def add_formatted_paragraph(document, reference, text=""):
    paragraph = document.add_paragraph()
    copy_paragraph_format(reference, paragraph)
    if text:
        run = paragraph.add_run(text)
        if reference.runs and reference.runs[0]._r.rPr is not None:
            run._r.insert(0, deepcopy(reference.runs[0]._r.rPr))
    return paragraph


def add_figure(document, image_reference, caption_reference, image_path, caption, width=5.04):
    paragraph = add_formatted_paragraph(document, image_reference)
    paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    add_formatted_paragraph(document, caption_reference, caption)


def remove_paragraphs_from(document, start_index):
    for paragraph in list(document.paragraphs[start_index:]):
        paragraph._element.getparent().remove(paragraph._element)


def build_report():
    document = Document(SOURCE_DOCX)

    title_reference = document.paragraphs[85]
    section_reference = document.paragraphs[86]
    body_reference = document.paragraphs[87]
    blank_reference = document.paragraphs[89]
    image_reference = document.paragraphs[92]
    caption_reference = document.paragraphs[93]

    document.paragraphs[1].text = "2026/07/31"
    date_run = document.paragraphs[1].runs[0]
    source_date_run = document.paragraphs[2].runs[0]
    if source_date_run._r.rPr is not None:
        if date_run._r.rPr is not None:
            date_run._r.remove(date_run._r.rPr)
        date_run._r.insert(0, deepcopy(source_date_run._r.rPr))

    remove_paragraphs_from(document, 85)

    def title(text):
        add_formatted_paragraph(document, title_reference, text)

    def section(text):
        add_formatted_paragraph(document, section_reference, text)

    def body(text):
        add_formatted_paragraph(document, body_reference, text)

    def blank():
        add_formatted_paragraph(document, blank_reference)

    title(
        "④  Optimization of a 3D-Printed Permanent Magnet Coupling Using "
        "Genetic Algorithm and Taguchi Method（2021）"
    )
    section("{Introduction}")
    body(
        "本論文は，永久磁石カップリングの伝達トルクを維持しながら構造体積を削減する形状最適化を対象とし，"
        "多目的遺伝的アルゴリズム（GA）とTaguchi法を有限要素解析および実機試験によって比較した研究である．"
        "永久磁石カップリングは非接触で動力を伝達できる一方，磁石寸法，ヨーク形状，空隙および相対角度に対して"
        "磁束とトルクが非線形に変化するため，単一寸法の経験的調整では材料使用量と性能を同時に改善しにくい．"
    )
    body(
        "著者らは，積層造形により複雑形状を製作できることを前提に，従来の直径や磁石寸法だけでなく，"
        "磁束密度の低い領域から材料を除去する局所形状を設計変数とした．この考え方は，任意形状を無制限に広げる"
        "のではなく，電磁場から感度の低い領域を抽出し，製作可能な形状自由度へ変換して探索する点に特徴がある．"
    )
    blank()
    section("{Coupling model and finite-element analysis}")
    body(
        "初期構造は，N52 Nd-Fe-B磁石を有する駆動側円板と，けい素鋼製の歯を有する従動側円板から成り，"
        "両者の空隙は1.5 mmである．外径30 mm，内径12.5 mm，円板厚1.2 mmを基準とし，"
        "相対偏角θに応じて従動側へ磁気トルクτを発生する．初期設計はSimcenter MAGNETで解析され，"
        "最大トルク73.0×10^-3 N m，構造体積2.31×10^-6 m^3を得た．最大トルクはθ=17 deg付近で生じる．"
    )
    add_figure(
        document,
        image_reference,
        caption_reference,
        FIGURE_DIR / "paper4_coupling_initial_design.jpg",
        "図10　永久磁石カップリングの初期構造と磁石極性（原著Fig. 1）",
        width=4.65,
    )
    body(
        "形状探索には，円板上の12点の局所厚さを表す12変数を用いる．各変数の移動範囲は0から1 mmであり，"
        "同一形状を円周方向へ反復して閉じた円板を構成する．目的は最大トルクτ_maxの増加と体積Vの減少であり，"
        "一般形はmin f(x)=[f_1(x),...,f_n(x)]^T，g_i(x)<=0，h_j(x)=0として表される．"
        "本研究ではτ_maxとVが解析式ではなくFEMから返されるため，各候補のCAD生成と磁場解析が目的関数評価になる．"
    )
    blank()
    section("{Multi-objective optimization}")
    body(
        "GAでは，個体を12変数の組として初期化し，FEMで目的値を計算した後，Pareto優越関係に基づいて順位と"
        "適応度を与える．近傍へ個体が集中することを防ぐfitness sharingを併用し，tournament selection，"
        "single-point crossoverおよびuniform mutationによって次世代を生成する．実装ではMatlabが候補を出力し，"
        "SolidWorksが3D形状を生成し，MAGNETが最大トルクを評価してMatlabへ返す閉ループを構成した．"
    )
    add_figure(
        document,
        image_reference,
        caption_reference,
        FIGURE_DIR / "paper4_coupling_optimization_flow.jpg",
        "図11　CAD生成，FEM解析および最適化を閉ループ化した処理手順（原著Fig. 5）",
        width=4.85,
    )
    body(
        "GAは50個体，30世代，合計1500回のFEM評価を行った．得られた設計は最大トルク74.5×10^-3 N m，"
        "体積1.91×10^-6 m^3であり，初期設計と比較してトルクを約2%増加させながら体積を約17%削減した．"
        "一方，計算時間は約50 hを要した．多目的探索はPareto解集合を得られるが，高価なFEMを候補ごとに呼び出すため，"
        "評価回数が計算時間へ直結することが定量的に示されている．"
    )
    body(
        "Taguchi法では，1変数を2水準，残る11変数を5水準とするL50直交表を用い，50通りのFEM評価から"
        "目的関数f=V/τを最小にする水準組合せを求めた．最適設計は最大トルク75.0×10^-3 N m，"
        "体積1.91×10^-6 m^3であり，GAとほぼ同一の形状と性能を1.5 hで得た．トルク密度は"
        "31.6から39.3 kN m/m^3へ増加し，20%以上改善した．ただし，Taguchi法は高次元かつ強く相互作用する"
        "多目的問題へ一般化しにくく，GAを常に置き換えるものではない．"
    )
    blank()
    section("{Manufacturing and experimental validation}")
    body(
        "初期形状とTaguchi最適形状は，6.5%けい素鋼を選択的レーザ溶融で造形し，空隙を1，1.5，2から3 mmへ"
        "設定できる試験台で評価した．相対角度は1 deg刻みで固定し，従動軸の腕に作用する力F_Rを秤で測定して"
        "τ=F_R rへ換算した．各設計を4回測定し，平均トルク曲線をFEM結果と比較している．"
    )
    add_figure(
        document,
        image_reference,
        caption_reference,
        FIGURE_DIR / "paper4_coupling_experimental_validation.jpg",
        "図12　初期・最適設計のトルク角特性に関するFEMと実測の比較（原著Fig. 15）",
        width=5.04,
    )
    body(
        "0から7 degでは解析曲線と実測曲線が近接したが，7から17 degでは実測トルクが解析値より低下した．"
        "著者らは，試験台構造，軸受および秤の分解能が差へ影響したと考察している．実測では最大トルクが"
        "初期71.0×10^-3 N m，最適72.0×10^-3 N mであり，トルク密度は30.7から37.7 kN m/m^3へ向上した．"
        "したがって，数値最適化だけでなく，製作誤差と支持機構を含む実測によって性能を再判定する必要がある．"
    )
    blank()
    section("{Conclusions}")
    body(
        "本論文は，永久磁石カップリングについて，設計変数の定義，CAD生成，FEM評価，多目的探索，積層造形，"
        "実測検証までを一つの流れとして示した．現在のリング型磁気結合研究に対しては，磁石数や形状自由度を"
        "増やすだけでなく，物理感度に基づく変数化，Pareto評価，高忠実度解析へ送る候補数の制御，および"
        "実測によるsim-to-real差の確認が必要であることを示す直接的な先行事例である．"
    )
    blank()
    blank()

    title(
        "⑤  A New Approximation for Calculating the Attraction Force in Cylindrical "
        "Permanent Magnets Arrays and Cylindrical Linear Single-Axis-Actuator（2019）"
    )
    section("{Introduction}")
    body(
        "本論文は，有限寸法の円柱永久磁石とその配列について，磁場および相互作用力を解析的に近似する手法を提案し，"
        "他手法と実測値によって検証した研究である．円柱磁石を点双極子として扱うと近接場の端面形状を失う一方，"
        "円形端面の厳密積分を多数磁石の全組合せへ適用すると計算量が大きい．著者らは円形端面を複数の矩形へ分割し，"
        "円柱全体を有限個の直方体磁石へ置換することで，精度を分割数Nにより調整できる近似を構成した．"
    )
    body(
        "対象は，同軸および横ずれを有する二円柱磁石，2×2，4×4，6×6の磁石配列，ならびに永久磁石と薄形コイルから"
        "成る単軸アクチュエータである．単体磁石の磁場から配列の合力までを同じ重ね合わせ原理で記述しており，"
        "多数の13 mm円盤磁石を配置する現在の磁気結合モデルへ直接適用できる．"
    )
    blank()
    section("{Coulombian model and finite-size decomposition}")
    body(
        "一様磁化Mを仮定すると，Coulombian modelでは磁石を表面磁荷密度σ*=M・nを有する二つの端面へ置換する．"
        "観測点Rにおける磁気スカラポテンシャルは，φ(R)=(1/4πμ_0)∫_S σ*/R_e dSで表され，"
        "磁束密度はその勾配から求められる．円形端面を内接矩形へ分割し，各矩形の解析式を加算することで，"
        "円柱磁石のB_x，B_y，B_zを有限和として評価する．"
    )
    add_figure(
        document,
        image_reference,
        caption_reference,
        FIGURE_DIR / "paper5_cuboid_decomposition.jpg",
        "図13　円柱永久磁石を有限個の直方体磁石へ分割する近似（原著Fig. 6）",
        width=4.4,
    )
    body(
        "分割数Nを増加させると矩形集合の外形が円へ近づき，磁場分布の幾何学誤差を小さくできる．"
        "直径1 mm，高さ10 mm，残留磁束密度1.15 Tの例では，N=4よりN=40の方が円柱周囲の磁場分布を滑らかに再現した．"
        "さらに，磁石軸上を移動するセンサで測定したB_zと計算結果を比較し，距離変化に対する整合を確認している．"
        "このためNは単なる計算設定ではなく，速度と有限寸法精度を交換する明示的な忠実度変数になる．"
    )
    blank()
    section("{Interaction force between cylindrical magnets}")
    body(
        "二円柱磁石をそれぞれN個およびM個の直方体要素へ分割し，要素sとtの相互作用エネルギーE_s,tを求める．"
        "力はF_s,t=grad(E_s,t)であり，全磁石力はF_magnet=Σ_sΣ_t F_s,tとなる．この式は軸方向力だけでなく"
        "三方向成分を返すため，中心からのx，y変位に対する復元力および局所的な非復元成分を同じ評価器で検査できる．"
    )
    body(
        "横方向へ移動する二磁石では，移動方向の力が変位符号に応じて反転し，対称位置で逆向きの力を生じる．"
        "軸方向距離を変えた試験では，計算曲線が実験挙動を良好に再現した．ただし，論文の検証条件は規則的な"
        "単体または小規模配列であり，任意姿勢，全磁石相互作用および極小空隙へ拡張する場合は，分割収束試験と"
        "高忠実度FEMによる再検証が必要である．"
    )
    blank()
    section("{Force of permanent-magnet arrays}")
    body(
        "著者らは，半径2 mm，高さ8 mmの円柱磁石を交互磁化して2×2，4×4および6×6の配列を構成し，"
        "対向配列間の吸引力を距離zの関数として算出した．磁石数の増加により近距離力は大きくなるが，"
        "いずれも距離とともに急減する．すなわち，総磁石数だけでは性能を決定できず，各磁石対の空隙，極性，"
        "横ずれおよび配置周期を含む全対相互作用を評価する必要がある．"
    )
    add_figure(
        document,
        image_reference,
        caption_reference,
        FIGURE_DIR / "paper5_array_force_validation.jpg",
        "図14　2×2，4×4および6×6円柱磁石配列の力距離特性（原著Fig. 15）",
        width=4.7,
    )
    body(
        "4×4配列の実測比較では，z=0，0.5，10，14.8 mmにおける力がそれぞれ9.550，5.110，0.320，"
        "0.100 Nであり，提案法は9.400，5.493，0.295，0.118 Nを与えた．絶対誤差は条件により変化するため，"
        "一点の一致だけでなく，使用空隙全域における誤差を確認すべきである．一方，距離減衰と磁石数依存を"
        "低計算費用で再現できることは，大規模探索のsurrogateとして有用である．"
    )
    blank()
    section("{Actuator validation and conclusions}")
    body(
        "同じ近似は，永久磁石と薄形コイルから成る円柱単軸アクチュエータにも拡張された．コイルを四つの直線導体へ"
        "置換し，ローレンツ力F=∫_V J_c×B dVを解析的に評価する．軸方向位置を変えたF_zは実測値と近接し，"
        "円柱磁石の有限寸法近似が磁石間力だけでなく磁石-導体相互作用にも適用できることを示した．"
    )
    add_figure(
        document,
        image_reference,
        caption_reference,
        FIGURE_DIR / "paper5_actuator_force_validation.jpg",
        "図15　単軸アクチュエータの軸方向力に関する近似計算と実測の比較（原著Fig. 19）",
        width=4.55,
    )
    body(
        "本論文の意義は，円柱磁石を有限個の解析可能な直方体へ分割し，分割数Nによって精度と計算時間を制御した点にある．"
        "現在のGPU探索では，粗いNで多数候補を選別し，上位候補だけNを増やして再評価し，最後にFEMと実測へ移す"
        "多忠実度構成が妥当である．この順序により，有限寸法を無視した点双極子近似だけで候補を確定する危険と，"
        "全候補を最高忠実度で解析する計算費用の双方を抑えられる．"
    )
    body(
        "台車用リングへ適用する際は，各候補について分割数Nを段階的に増加させ，評価対象であるF_x，F_yおよびτ_zの"
        "相対変化が所定値以下になるまで収束を確認する必要がある．その上で，動作範囲内の(x,y,θ)格子へ全磁石対の力を"
        "重ね合わせ，q=[x,y,l_ref θ]^Tに対する復元仕事-q^T f(q)>0，最小空隙g_min>0，"
        "および局所剛性行列の最小固有値λ_min>0を同時に検査する．"
    )
    body(
        "ただし，原著の力検証は主として吸引配置と軸方向変位を対象としている．本研究の同極反発配置では磁荷符号を"
        "反転できるものの，横ずれ，回転，磁石個体差および治具による磁化方向誤差を含む実測妥当性は原著だけでは保証"
        "されない．したがって，13 mm磁石の実測力マップで残留磁束密度と寸法公差を同定し，FEM，Monte Carloばらつき解析，"
        "静的三自由度試験，10 kg台車動的試験の順に検証することが必要である．"
    )
    body(
        "以上より，本手法は最終的な実機保証そのものではなく，大規模探索と高忠実度検証を接続する解析surrogateとして"
        "位置付けるべきである．解析近似，FEMおよび実測の三段階で同じ評価点を照合することにより，計算速度だけでなく，"
        "非接触余裕と三自由度復元性を根拠として候補を選別できる．"
    )

    document.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    build_report()
